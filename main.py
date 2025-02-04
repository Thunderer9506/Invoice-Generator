import subprocess
import shutil
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk
import docx
import num2words

class InvoiceAutomation:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Invoice Automation")
        self.root.geometry("600x600")

        self.upperFrame = ttk.Frame(self.root)
        self.upperFrame.place(relx=0, rely=0, relwidth=1, relheight=0.2)
        self.upperFrame.columnconfigure((0,1,2,3), weight=1)
        self.upperFrame.rowconfigure((0,1,2), weight=1)

        self.midFrame = ttk.Frame(self.root)
        self.midFrame.place(relx=0, rely=0.3, relwidth=1, relheight=0.4)
        self.midFrame.columnconfigure((0,1,2,3), weight=1)
        self.midFrame.rowconfigure((0,1,2,3,4,5), weight=1)

        self.endFrame = ttk.Frame(self.root)
        self.endFrame.place(relx=0, rely=0.7, relwidth=1, relheight=0.1)
        self.endFrame.columnconfigure((0,1,2), weight=1)
        self.endFrame.rowconfigure((0,1), weight=1)

        #Lables
        self.date_label = ttk.Label(self.upperFrame, text="Date")
        self.invoice_label = ttk.Label(self.upperFrame, text="Invoice Number")

        self.clientName_label = ttk.Label(self.upperFrame, text="Client Name")
        self.clientAddress_label = ttk.Label(self.upperFrame, text="Client Address")
        self.clientGST_label = ttk.Label(self.upperFrame, text="Client GST")
        
        self.description_label = tk.Label(self.midFrame, text="Description")
        self.quantity_label = tk.Label(self.midFrame, text="Quantity")
        self.rate_label = tk.Label(self.midFrame, text="Rate")
        self.amount_label = tk.Label(self.midFrame, text="Amount")
        
        self.igst_label = ttk.Label(self.endFrame, text="IGST")
        self.sgst_label = ttk.Label(self.endFrame, text="SGST")
        self.cgst_label = ttk.Label(self.endFrame, text="CGST")

        #Entry Widget 
        self.date_entry = ttk.Entry(self.upperFrame,width=30)
        self.invoice_entry = ttk.Entry(self.upperFrame,width=30)

        self.clientName_entry = ttk.Entry(self.upperFrame,width=30)
        self.clientAddress_entry = ttk.Entry(self.upperFrame,width=30)
        self.clientGST_entry = ttk.Entry(self.upperFrame,width=30)

        self.description1_entry = ttk.Entry(self.midFrame,width=50)
        self.description2_entry = ttk.Entry(self.midFrame,width=50)
        self.description3_entry = ttk.Entry(self.midFrame,width=50)
        self.description4_entry = ttk.Entry(self.midFrame,width=50)
        self.description5_entry = ttk.Entry(self.midFrame,width=50)

        self.quantity1_entry = ttk.Entry(self.midFrame,width=10)
        self.quantity2_entry = ttk.Entry(self.midFrame,width=10)
        self.quantity3_entry = ttk.Entry(self.midFrame,width=10)
        self.quantity4_entry = ttk.Entry(self.midFrame,width=10)
        self.quantity5_entry = ttk.Entry(self.midFrame,width=10)

        self.rate1_entry = ttk.Entry(self.midFrame,width=10)
        self.rate2_entry = ttk.Entry(self.midFrame,width=10)
        self.rate3_entry = ttk.Entry(self.midFrame,width=10)
        self.rate4_entry = ttk.Entry(self.midFrame,width=10)
        self.rate5_entry = ttk.Entry(self.midFrame,width=10)

        self.amount1_entry = ttk.Entry(self.midFrame,width=20)
        self.amount2_entry = ttk.Entry(self.midFrame,width=20)
        self.amount3_entry = ttk.Entry(self.midFrame,width=20)
        self.amount4_entry = ttk.Entry(self.midFrame,width=20)
        self.amount5_entry = ttk.Entry(self.midFrame,width=20)

        self.igst_var = tk.BooleanVar(value=False)
        self.sgst_var = tk.BooleanVar(value=False)
        self.cgst_var = tk.BooleanVar(value=False)

        self.igst_checkbutton = ttk.Checkbutton(self.endFrame,text="IGST",variable=self.igst_var,onvalue=True,offvalue=False)
        self.sgst_checkbutton = ttk.Checkbutton(self.endFrame,text="SGST",variable=self.sgst_var,onvalue=True,offvalue=False)
        self.cgst_checkbutton = ttk.Checkbutton(self.endFrame,text="CGST",variable=self.cgst_var,onvalue=True,offvalue=False)
        

        self.create_Button = ttk.Button(self.root, text="Create Invoice", command=self.create_invoice)

        # Now pack every widget
        # Create a canvas and a scrollbar
        

        # Now use grid to place widgets in the scrollable frame
        self.date_label.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.date_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        self.invoice_label.grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.invoice_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        self.clientName_label.grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.clientName_entry.grid(row=0, column=3, padx=5, pady=5, sticky="ew")

        self.clientAddress_label.grid(row=1, column=2, padx=5, pady=5, sticky="w")
        self.clientAddress_entry.grid(row=1, column=3, padx=5, pady=5, sticky="ew")

        self.clientGST_label.grid(row=2, column=2, padx=5, pady=5, sticky="w")
        self.clientGST_entry.grid(row=2, column=3, padx=5, pady=5, sticky="ew")

        self.description_label.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.quantity_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        self.rate_label.grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.amount_label.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        self.description1_entry.grid(row=1, column=0, padx=5, pady=5)
        self.description2_entry.grid(row=2, column=0, padx=5, pady=5)
        self.description3_entry.grid(row=3, column=0, padx=5, pady=5)
        self.description4_entry.grid(row=4, column=0, padx=5, pady=5)
        self.description5_entry.grid(row=5, column=0, padx=5, pady=5)

        self.quantity1_entry.grid(row=1, column=1, padx=5, pady=5)
        self.quantity2_entry.grid(row=2, column=1, padx=5, pady=5)
        self.quantity3_entry.grid(row=3, column=1, padx=5, pady=5)
        self.quantity4_entry.grid(row=4, column=1, padx=5, pady=5)
        self.quantity5_entry.grid(row=5, column=1, padx=5, pady=5)

        self.rate1_entry.grid(row=1, column=2, padx=5, pady=5)
        self.rate2_entry.grid(row=2, column=2, padx=5, pady=5)
        self.rate3_entry.grid(row=3, column=2, padx=5, pady=5)
        self.rate4_entry.grid(row=4, column=2, padx=5, pady=5)
        self.rate5_entry.grid(row=5, column=2, padx=5, pady=5)

        self.amount1_entry.grid(row=1, column=3, padx=5, pady=5)
        self.amount2_entry.grid(row=2, column=3, padx=5, pady=5)
        self.amount3_entry.grid(row=3, column=3, padx=5, pady=5)
        self.amount4_entry.grid(row=4, column=3, padx=5, pady=5)
        self.amount5_entry.grid(row=5, column=3, padx=5, pady=5)

        self.igst_label.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.igst_checkbutton.grid(row=1, column=0, padx=5, pady=5, sticky="ew")

        self.sgst_label.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        self.sgst_checkbutton.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        self.cgst_label.grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.cgst_checkbutton.grid(row=1, column=2, padx=5, pady=5, sticky="ew")

        self.create_Button.place(relx=0, rely=0.9, relwidth=1, relheight=0.1)

        self.root.mainloop()

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text,new_text)
    
    def subtotal(self):
        Amount1 = float(self.amount1_entry.get()) if self.amount1_entry.get() else 0
        Amount2 = float(self.amount2_entry.get()) if self.amount2_entry.get() else 0
        Amount3 = float(self.amount3_entry.get()) if self.amount3_entry.get() else 0
        Amount4 = float(self.amount4_entry.get()) if self.amount4_entry.get() else 0
        Amount5 = float(self.amount5_entry.get()) if self.amount5_entry.get() else 0
        sum = Amount1 + Amount2 + Amount3 + Amount4 + Amount5
        return sum
    
    def igst(self, amount):
        return float("%.2f"%(amount * 0.18))

    def cgst(self, amount):
        return float("%.2f"%(amount * 0.09))

    def sgst(self, amount):
        return float("%.2f"%(amount * 0.09))

    def total_in_words(self, total):
        rounded_total = round(total)
        total_words = num2words.num2words(rounded_total).title()
        if total == rounded_total:
            return f"Total (Round Off):\n{total_words} Only"
        else:
            return f"Total:\n {total_words}"

    def create_invoice(self):
        doc = docx.Document("template.docx")

        subtotal = self.subtotal()
        igst = self.igst(subtotal) if self.igst_var.get() else 0
        sgst = self.sgst(subtotal) if self.sgst_var.get() else 0
        cgst = self.cgst(subtotal) if self.cgst_var.get() else 0
        total = subtotal + igst + sgst + cgst
        total = "%.2f" % total
        
        try:
            replacements = {
                "[Date]": self.date_entry.get(),
                "[Invoice]": self.invoice_entry.get(),

                "[clientName]": self.clientName_entry.get().title(),
                "[clientAddress]": self.clientAddress_entry.get(),
                "[clientGST]": self.clientGST_entry.get(),

                "[Description1]": self.description1_entry.get() if self.description1_entry.get() else "",
                "[Description2]": self.description2_entry.get() if self.description2_entry.get() else "",
                "[Description3]": self.description3_entry.get() if self.description3_entry.get() else "",
                "[Description4]": self.description4_entry.get() if self.description4_entry.get() else "",
                "[Description5]": self.description5_entry.get() if self.description5_entry.get() else "",
                "[Rate1]": self.rate1_entry.get() if self.rate1_entry.get() else "",
                "[Rate2]": self.rate2_entry.get() if self.rate2_entry.get() else "",
                "[Rate3]": self.rate3_entry.get() if self.rate3_entry.get() else "",
                "[Rate4]": self.rate4_entry.get() if self.rate4_entry.get() else "",
                "[Rate5]": self.rate5_entry.get() if self.rate5_entry.get() else "",
                "[Quantity1]": self.quantity1_entry.get() if self.quantity1_entry.get() else "",
                "[Quantity2]": self.quantity2_entry.get() if self.quantity2_entry.get() else "",
                "[Quantity3]": self.quantity3_entry.get() if self.quantity3_entry.get() else "",
                "[Quantity4]": self.quantity4_entry.get() if self.quantity4_entry.get() else "",
                "[Quantity5]": self.quantity5_entry.get() if self.quantity5_entry.get() else "",
                "[Amount1]": self.amount1_entry.get() if self.amount1_entry.get() else "",
                "[Amount2]": self.amount2_entry.get() if self.amount2_entry.get() else "",
                "[Amount3]": self.amount3_entry.get() if self.amount3_entry.get() else "",
                "[Amount4]": self.amount4_entry.get() if self.amount4_entry.get() else "",
                "[Amount5]": self.amount5_entry.get() if self.amount5_entry.get() else "",
                "[subTotal]": str(subtotal),
                "[IGST]":str(igst),
                "[SGST]":str(sgst),
                "[CGST]":str(cgst),
                "[Total]":str(total),
                "[amountInword]":self.total_in_words(float(total)),
            }
        except ValueError:
            messagebox.showerror(title='Error',message="Invalid amount or price")
            return
        
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph,old_text,new_text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph,old_text,new_text)
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[('PDF documents', '*.pdf')])
        if not save_path:
            messagebox.showerror(title="Error", message="No file name provided")
            return

        doc.save('filled.docx')

        try:
            subprocess.run(args=[r'C:\Libre Office\program\soffice.exe', '--headless', '--convert-to', 'pdf', 'filled.docx', '--outdir', '.'], check=True)
            shutil.move("filled.pdf", save_path)
            messagebox.showinfo(title="success", message=f"Invoice created and saved successfully as {save_path}")
        except subprocess.CalledProcessError as e:
            messagebox.showerror(title="Error", message=f"Failed to convert DOCX to PDF: {e}")
        except PermissionError as e:
            messagebox.showerror(title="Error", message=f"Permission denied: {e}")
        except OSError as e:
            messagebox.showerror(title="Error", message=f"Failed to move the file: {e}")


if __name__ == "__main__":
    InvoiceAutomation()